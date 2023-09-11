import openai
from docx import Document

used_model = "gpt-3.5-turbo-16k"

def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)
    return transcription['text']


def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model=used_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points. pleas answer me in CHINESE / 中文"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model=used_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about. pleas answer me in CHINESE / 中文"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']



def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model=used_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible. pleas answer me in CHINESE / 中文"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model=used_model,
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely. pleas answer me in CHINESE / 中文"
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

# GENERATE ONE
# audio_file_path = "voice_clips/xxx.wav"
# print("working in ===transcribe_audio(audio_file_path)===")
# transcription = transcribe_audio(audio_file_path)
# print("working in ===meeting_minutes(transcription)===")
# minutes = meeting_minutes(transcription)
# # print(minutes)
# save_as_docx(minutes, 'xxx.docx')


def save_ori_trans_docx(transcription, docx_name):
    # 创建一个新的Word文档
    doc = Document()
    # 添加transcription到文档中
    doc.add_paragraph(transcription)
    # 保存文档
    doc.save( 'meetings_origin/'+'trans_'+ docx_name+'.docx')

# GENERATE MULTIPLE
import os
def voice2docx(audio_file_path, docx_name):
    print("working for", docx_name)
    print("working in ===transcribe_audio(audio_file_path)===")
    transcription = transcribe_audio(audio_file_path)
    save_ori_trans_docx(transcription=transcription, docx_name=docx_name)
    print("working in ===meeting_minutes(transcription)===")
    minutes = meeting_minutes(transcription)
    # print(minutes)
    save_as_docx(minutes, 'meetings_summary_clips/' + docx_name+'.docx')


def traverse_files(directory):
    for dirpath, dirnames, filenames in os.walk(directory):
        for filename in filenames:
            if filename.endswith('.wav'):
                voice2docx(audio_file_path=os.path.join(dirpath, filename),
                           docx_name=filename.strip('.wav'))

traverse_files('voice_clips')